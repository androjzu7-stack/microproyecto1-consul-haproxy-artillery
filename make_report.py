# make_report.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

OUT = "Microproyecto1_Entrega_Final.docx"
EVI = "evidencias"

FIGS = [
    ("01_vagrant_status.png", "Vagrant status mostrando lb, web1, web2 y client en estado running."),
    ("02_consul_members.png", "Salida de consul members mostrando los nodos lb, web1 y web2 en estado alive."),
    ("03_dns_srv.png", "Respuesta DNS SRV del servicio web (descubrimiento de instancias)."),
    ("04_consul_ui_services.png", "Consul UI (Services) mostrando servicio web con 8 instancias."),
    ("05_scripts_tests.png", "Listado de scripts y archivos de pruebas Artillery (baseline/spike/soak + results)."),
    ("06_vagrantfile.png", "Vagrantfile con definición de lb/web1/web2/client y ejecución de provisioners."),
    ("07_ping_lb.png", "Ping a 192.168.100.10 (lb) desde el host."),
    ("08_curl_ip_200.png", "Respuesta HTTP 200 OK accediendo al servicio por IP del balanceador."),
    ("09_haproxy_stats_ip.png", "HAProxy Stats accesible por IP (192.168.100.10:1936) mostrando backends UP."),
    ("10_balanceo_8.png", "Distribución uniforme hacia 8 instancias (web1/web2 y puertos 3000–3004)."),
    ("11_503_personalizado.png", "Página 503 personalizada mostrada cuando no hay backends disponibles (acceso por IP)."),
    ("12_artillery_baseline.png", "Artillery baseline ejecutado desde la VM client (resumen final)."),
    ("13_artillery_spike.png", "Artillery spike (pico) — resumen final del escenario."),
    ("14_artillery_soak.png", "Artillery soak (carga sostenida) con HAProxy Stats en paralelo."),
    ("15_stick_table_client.png", "show table http_front mostrando key=192.168.100.13 (VM client) en stick-table."),
]

def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    tcPr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

def add_fig(doc, n, filename, caption):
    path = os.path.join(EVI, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.7))
        p = doc.add_paragraph(f"Figura {n}. {caption}")
        p.runs[0].italic = True
    else:
        p = doc.add_paragraph(f"[FALTA] Figura {n}: {filename}")
        p.runs[0].italic = True

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Portada
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MICROPROYECTO 1\n"); r.bold = True; r.font.size = Pt(22)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cluster Consul + Balanceador HAProxy + Artillery en Vagrant"); r.bold = True; r.font.size = Pt(14)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Estudiante: Andres\n").bold = True
p.add_run("Fecha: Febrero 2026\n").bold = True
p.add_run("Provider: VMware Desktop | VMs: Ubuntu 22.04\n").bold = True

doc.add_page_break()

# Secciones
doc.add_heading("1. Objetivo", level=1)
doc.add_paragraph(
    "Implementar un clúster de HashiCorp Consul para service discovery con agentes en máquinas virtuales. "
    "Al menos dos nodos alojan una aplicación web en Node.js. Las peticiones se realizan únicamente al balanceador "
    "HAProxy, el cual distribuye el tráfico hacia los servidores disponibles. Adicionalmente: GUI de HAProxy accesible "
    "desde el host, escalabilidad con réplicas, página 503 personalizada cuando no existan servidores disponibles y "
    "pruebas de carga con Artillery."
)

doc.add_heading("2. Arquitectura y topología", level=1)
doc.add_paragraph("Accesos desde el host (por IP del balanceador):")
for item in [
    "Servicio: http://192.168.100.10/",
    "HAProxy Stats: http://192.168.100.10:1936/",
    "Consul UI: http://192.168.100.10:8500/",
]:
    doc.add_paragraph(item, style=doc.styles["List Bullet"])

doc.add_heading("3. PREGUNTA 1 — Implementación de Cluster Consul", level=1)
code_block(doc, 'vagrant ssh lb -c "consul members"\n'
                'vagrant ssh lb -c "dig +short SRV _web._tcp.service.consul @127.0.0.1 -p 8600"')

doc.add_heading("4. PREGUNTA 2 — Aprovisionamiento (Vagrant)", level=1)
code_block(doc, "vagrant up\nvagrant provision\nls -l scripts\nls -l tests")

doc.add_heading("5. PREGUNTA 3 — Disponibilidad, balanceo y pruebas de carga", level=1)
doc.add_paragraph("5.1 Acceso por IP y GUI de HAProxy")
code_block(doc, "ping -c 3 192.168.100.10\ncurl -i http://192.168.100.10 | head -n 10")

doc.add_paragraph("5.2 Balanceo y escalabilidad (8 instancias)")
code_block(doc, "for i in {1..40}; do curl -s http://192.168.100.10 | jq -r '.node+\":\"+(.port|tostring)'; done | sort | uniq -c")

doc.add_paragraph("5.3 Caso sin servidores (503 personalizado)")
code_block(doc,
          "vagrant ssh web1 -c \"sudo systemctl stop --all 'web@*' && sudo consul reload\"\n"
          "vagrant ssh web2 -c \"sudo systemctl stop --all 'web@*' && sudo consul reload\"\n"
          "Abrir: http://192.168.100.10/")

doc.add_paragraph("5.4 Pruebas de carga con Artillery (VM client)")
code_block(doc,
          "vagrant ssh client -c \"artillery run ~/tests/baseline.yml\"\n"
          "vagrant ssh client -c \"artillery run ~/tests/spike.yml\"\n"
          "vagrant ssh client -c \"artillery run ~/tests/soak.yml\"")

# Tabla resumida (ajústala si quieres)
doc.add_paragraph("Resumen (desde capturas):")
t = doc.add_table(rows=1, cols=6); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text="Escenario"; hdr[1].text="http.responses"; hdr[2].text="vusers.failed"
hdr[3].text="Req/s"; hdr[4].text="p95 resp (s)"; hdr[5].text="p99 resp (s)"
rows = [
    ("Baseline", "600", "0", "20/s", "4", "7"),
    ("Spike", "1000", "0", "28/s", "2", "5"),
    ("Soak", "300", "0", "30/s", "2", "3"),
]
for r in rows:
    row = t.add_row().cells
    for i,v in enumerate(r): row[i].text = str(v)

doc.add_paragraph("5.5 Cliente visible en frontend (stick-table)")
code_block(doc, "vagrant ssh lb -c \"echo 'show table http_front' | sudo socat stdio /run/haproxy/admin.sock\"")

# Figuras
doc.add_heading("6. Evidencias", level=1)
fig_no = 1
for fn, cap in FIGS:
    add_fig(doc, fig_no, fn, cap)
    fig_no += 1

doc.add_heading("7. Conclusiones", level=1)
doc.add_paragraph(
    "Se implementó Consul (service discovery), HAProxy (balanceo + GUI por IP), escalabilidad con 8 instancias, "
    "manejo de indisponibilidad con 503 personalizado y pruebas de carga con Artillery desde una VM client. "
    "Los escenarios se ejecutaron sin fallos."
)

doc.add_heading("8. Referencias", level=1)
for ref in [
    "HashiCorp Consul — Documentación oficial.",
    "HAProxy — Documentación oficial.",
    "Artillery — Documentación oficial.",
]:
    doc.add_paragraph(ref, style=doc.styles["List Bullet"])

doc.save(OUT)
print("OK ->", OUT)